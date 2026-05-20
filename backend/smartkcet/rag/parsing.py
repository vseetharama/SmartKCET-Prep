"""Document parsing and OCR helpers.

These utilities were lifted verbatim from ``backend/app.py`` and grouped
here so the upload pipeline can be evolved (per-subject isolation, OCR
warning aggregation per design.md §5) without touching unrelated modules.
"""

from __future__ import annotations

import io
import logging
from typing import List

import cv2
import fitz
import numpy as np
import pytesseract
from PIL import Image
from docx import Document as DocxDocument

logger = logging.getLogger("smartkcet.rag.parsing")


def preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """Apply denoising and adaptive thresholding for better OCR accuracy."""

    try:
        img_np = np.array(img.convert("RGB"))
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        gray = cv2.fastNlMeansDenoising(gray, h=10)
        thresh = cv2.adaptiveThreshold(
            gray,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            10,
        )
        return Image.fromarray(thresh)
    except Exception as exc:  # pragma: no cover - defensive logging
        print(f"OCR preprocessing failed: {exc}")
        return img


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF; fall back to OCR for pages with little text."""

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    logger.info(
        "PDF opened: %d page(s), metadata=%r",
        len(doc),
        doc.metadata.get("title", ""),
    )
    pages: List[str] = []
    for page_num, page in enumerate(doc):
        text = page.get_text().strip()
        if len(text) > 50:
            logger.debug(
                "Page %d: extracted %d chars via text layer (direct extraction)",
                page_num + 1,
                len(text),
            )
            pages.append(text)
            continue

        # Text layer is empty or too short — attempt OCR fallback
        # But first, keep any short text we did find
        if text:
            logger.info(
                "Page %d: text layer has only %d chars (below 50-char threshold), "
                "keeping it and also attempting OCR",
                page_num + 1,
                len(text),
            )
            pages.append(text)
        else:
            logger.info(
                "Page %d: text layer is empty, attempting OCR fallback",
                page_num + 1,
            )
        try:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = preprocess_for_ocr(img)
            try:
                ocr_text = pytesseract.image_to_string(
                    img, lang="eng", config="--psm 6 --oem 3"
                )
                if ocr_text.strip():
                    logger.info(
                        "Page %d: OCR produced %d chars",
                        page_num + 1,
                        len(ocr_text.strip()),
                    )
                    pages.append(ocr_text)
                else:
                    logger.warning(
                        "Page %d: OCR returned empty text (truly unreadable page)",
                        page_num + 1,
                    )
            except Exception as ocr_exc:
                logger.warning(
                    "Page %d: OCR failed: %s (tesseract may not be installed or configured)",
                    page_num + 1,
                    ocr_exc,
                )
        except Exception as exc:
            logger.warning(
                "Page %d: pixmap/OCR pipeline failed: %s",
                page_num + 1,
                exc,
            )

    total_text = "\n".join(pages)
    logger.info(
        "PDF extraction complete: %d page(s) yielded text, total %d chars",
        len(pages),
        len(total_text),
    )
    return total_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    logger.info("DOCX extraction: %d chars from %d paragraphs", len(text), len(doc.paragraphs))
    return text


def extract_text_from_txt(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    logger.info("TXT extraction: %d chars", len(text))
    return text


def chunk_text(text: str, size: int = 400, overlap: int = 80) -> List[str]:
    """Split ``text`` into overlapping word-windows ready for embedding."""

    words = text.split()
    chunks: List[str] = []
    i = 0
    while i < len(words):
        chunks.append(" ".join(words[i : i + size]))
        i += size - overlap
    return [c for c in chunks if len(c.strip()) > 30]
